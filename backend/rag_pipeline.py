"""
RAG Pipeline: Document processing, embedding, storage, and retrieval
Supports: PDF, DOCX, RTF, PNG, TXT
"""
import os
from pathlib import Path
from typing import List, Dict
import logging

# Document loaders
import pdfplumber
from docx import Document as DocxDocument
import pypandoc

# LangChain components
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain.chains import ConversationalRetrievalChain
from langchain.schema import Document

# Memory
from memory import get_memory

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration
VECTORSTORE_PATH = "data/vectorstore"
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

# Initialize Gemini
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    logger.warning("GOOGLE_API_KEY not found in environment variables")

# Global vectorstore and chain
_vectorstore = None
_qa_chain = None

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfplumber"""
    try:
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        logger.info(f"Extracted {len(text)} characters from PDF: {Path(file_path).name}")
        return text
    except Exception as e:
        logger.error(f"Error extracting PDF {file_path}: {str(e)}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx"""
    try:
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        logger.info(f"Extracted {len(text)} characters from DOCX: {Path(file_path).name}")
        return text
    except Exception as e:
        logger.error(f"Error extracting DOCX {file_path}: {str(e)}")
        return ""

def extract_text_from_rtf(file_path: str) -> str:
    """Extract text from RTF using pypandoc"""
    try:
        text = pypandoc.convert_file(file_path, 'plain', format='rtf')
        logger.info(f"Extracted {len(text)} characters from RTF: {Path(file_path).name}")
        return text
    except Exception as e:
        logger.error(f"Error extracting RTF {file_path}: {str(e)}")
        return ""

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        logger.info(f"Extracted {len(text)} characters from TXT: {Path(file_path).name}")
        return text
    except Exception as e:
        logger.error(f"Error extracting TXT {file_path}: {str(e)}")
        return ""

def extract_text_from_file(file_path: str) -> tuple:
    """
    Route to appropriate extractor based on file extension
    
    Returns:
        tuple: (text_content, filename)
    """
    file_path = Path(file_path)
    extension = file_path.suffix.lower()
    
    extractors = {
        '.pdf': extract_text_from_pdf,
        '.docx': extract_text_from_docx,
        '.rtf': extract_text_from_rtf,
        '.txt': extract_text_from_txt,
    }
    
    extractor = extractors.get(extension)
    if extractor:
        text = extractor(str(file_path))
        return text, file_path.name
    else:
        logger.warning(f"No extractor for file type: {extension}")
        return "", file_path.name

def process_documents(file_paths: List[str]) -> Dict:
    """
    Main document processing pipeline:
    1. Extract text from all documents
    2. Chunk the text
    3. Create embeddings
    4. Store in vector database
    
    Args:
        file_paths: List of absolute paths to documents
        
    Returns:
        Dict with processing statistics
    """
    global _vectorstore, _qa_chain
    
    try:
        # Extract text from all documents
        documents = []
        for file_path in file_paths:
            text, filename = extract_text_from_file(file_path)
            if text.strip():
                documents.append(
                    Document(
                        page_content=text,
                        metadata={"source": filename}
                    )
                )
        
        if not documents:
            raise ValueError("No text content extracted from uploaded files")
        
        # Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            length_function=len,
        )
        chunks = text_splitter.split_documents(documents)
        logger.info(f"Created {len(chunks)} text chunks")
        
        # Create embeddings and vector store
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=GOOGLE_API_KEY
        )
        
        # If vectorstore exists, add to it; otherwise create new
        if _vectorstore is not None:
            _vectorstore.add_documents(chunks)
            logger.info(f"Added {len(chunks)} chunks to existing vectorstore")
        else:
            _vectorstore = Chroma.from_documents(
                documents=chunks,
                embedding=embeddings,
                persist_directory=VECTORSTORE_PATH
            )
            logger.info(f"Created new vectorstore with {len(chunks)} chunks")
        
        # Initialize QA chain with memory
        _qa_chain = create_qa_chain(_vectorstore)
        
        return {
            "status": "success",
            "documents_processed": len(documents),
            "chunks_count": len(chunks)
        }
    
    except Exception as e:
        logger.error(f"Error processing documents: {str(e)}")
        raise

def create_qa_chain(vectorstore):
    """
    Create a conversational retrieval chain with Gemini and memory
    
    Args:
        vectorstore: Chroma vectorstore instance
        
    Returns:
        ConversationalRetrievalChain
    """
    llm = ChatGoogleGenerativeAI(
        model="gemini-1.5-pro",
        google_api_key=GOOGLE_API_KEY,
        temperature=0.7,
        convert_system_message_to_human=True
    )
    
    memory = get_memory()
    
    qa_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(search_kwargs={"k": 4}),
        memory=memory,
        return_source_documents=True,
        verbose=True
    )
    
    return qa_chain

def query_rag(question: str) -> Dict:
    """
    Query the RAG system with conversation memory
    
    Args:
        question: User's question
        
    Returns:
        Dict with answer and sources
    """
    global _qa_chain, _vectorstore
    
    try:
        # Initialize vectorstore if not exists
        if _vectorstore is None:
            # Try to load existing vectorstore
            if Path(VECTORSTORE_PATH).exists():
                embeddings = GoogleGenerativeAIEmbeddings(
                    model="models/embedding-001",
                    google_api_key=GOOGLE_API_KEY
                )
                _vectorstore = Chroma(
                    persist_directory=VECTORSTORE_PATH,
                    embedding_function=embeddings
                )
                _qa_chain = create_qa_chain(_vectorstore)
                logger.info("Loaded existing vectorstore")
            else:
                return {
                    "answer": "No documents have been uploaded yet. Please upload documents first.",
                    "sources": []
                }
        
        # Query the chain
        result = _qa_chain({"question": question})
        
        # Extract sources
        sources = [doc.metadata.get("source", "Unknown") 
                  for doc in result.get("source_documents", [])]
        
        return {
            "answer": result.get("answer", ""),
            "sources": list(set(sources))  # Remove duplicates
        }
    
    except Exception as e:
        logger.error(f"Error querying RAG: {str(e)}")
        return {
            "answer": f"An error occurred: {str(e)}",
            "sources": []
        }

def reset_vectorstore():
    """Clear the vectorstore and start fresh"""
    global _vectorstore, _qa_chain
    
    _vectorstore = None
    _qa_chain = None
    
    # Remove persisted data
    import shutil
    if Path(VECTORSTORE_PATH).exists():
        shutil.rmtree(VECTORSTORE_PATH)
        logger.info("Vectorstore cleared")
